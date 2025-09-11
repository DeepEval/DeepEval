from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


class DocFileHandler:
    """
    This is a class that handles Word documents and provides functionalities for reading, writing, and modifying the content of Word documents.
    """

    def __init__(self, file_path):
        """
        Initializes the DocFileHandler object with the specified file path.
        :param file_path: str, the path to the Word document file.
        """
        self.file_path = file_path
        self.document = Document(file_path) if file_path else Document()

    def read_text(self):
        """
        Reads the content of a Word document and returns it as a string.
        :return: str, the content of the Word document.
        """
        full_text = []
        for paragraph in self.document.paragraphs:
            full_text.append(paragraph.text)
        return '\n'.join(full_text)

    def write_text(self, content, font_size=12, alignment='left'):
        """
        Writes the specified content to a Word document.
        :param content: str, the text content to write.
        :param font_size: int, optional, the font size of the text (default is 12).
        :param alignment: str, optional, the alignment of the text ('left', 'center', or 'right'; default is 'left').
        :return: bool, True if the write operation is successful, False otherwise.
        """
        paragraph = self.document.add_paragraph()
        run = paragraph.add_run(content)
        run.font.size = Pt(font_size)
        self._set_paragraph_alignment(paragraph, alignment)
        return True

    def add_heading(self, heading, level=1):
        """
        Adds a heading to the Word document.
        :param heading: str, the text of the heading.
        :param level: int, optional, the level of the heading (1, 2, 3, etc.; default is 1).
        :return: bool, True if the heading is successfully added, False otherwise.
        """
        paragraph = self.document.add_paragraph(heading, style=self.document.styles['Heading ' + str(level)])
        return True

    def add_table(self, data):
        """
        Adds a table to the Word document with the specified data.
        :param data: list of lists, the data to populate the table.
        :return: bool, True if the table is successfully added, False otherwise.
        """
        table = self.document.add_table(rows=len(data), cols=len(data[0]))
        for row_index, row in enumerate(data):
            for col_index, cell_value in enumerate(row):
                table.cell(row_index, col_index).text = str(cell_value)
        return True

    def _get_alignment_value(self, alignment):
        """
        Returns the alignment value corresponding to the given alignment string.
        :param alignment: str, the alignment string ('left', 'center', or 'right').
        :return: int, the alignment value.
        """
        if alignment == 'left':
            return WD_PARAGRAPH_ALIGNMENT.LEFT
        elif alignment == 'center':
            return WD_PARAGRAPH_ALIGNMENT.CENTER
        elif alignment == 'right':
            return WD_PARAGRAPH_ALIGNMENT.RIGHT
        else:
            raise ValueError("Invalid alignment value. Choose from 'left', 'center', or 'right'.")

    def _set_paragraph_alignment(self, paragraph, alignment):
        """
        Sets the alignment of a paragraph.
        :param paragraph: docx.paragraph.Paragraph, the paragraph to align.
        :param alignment: str, the alignment string ('left', 'center', or 'right').
        """
        paragraph.alignment = self._get_alignment_value(alignment)



if __name__ == "__main__":
    instance = DocFileHandler("test.docx")  # Replace "test.docx" with your file path

    # Test case for read_text
    output = instance.read_text()
    print("Read Text:\n", output)

    # Test case for write_text
    instance.write_text("This is some new text.", font_size=14, alignment='center')
    instance.document.save("test.docx")  

    # Test case for add_heading
    instance.add_heading("This is a Heading", level=2)
    instance.document.save("test.docx")

    # Test case for add_table
    data = [
        ["Name", "Age", "City"],
        ["John Doe", 30, "New York"],
        ["Jane Doe", 25, "London"]
    ]
    instance.add_table(data)
    instance.document.save("test.docx")