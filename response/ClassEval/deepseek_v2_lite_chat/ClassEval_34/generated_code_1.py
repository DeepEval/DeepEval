from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

class DocFileHandler:
    """
    This is a class that handles Word documents and provides functionalities for reading, writing, and modifying the content of Word documents.
    """

    def __init__(self, file_path):
        """
        Initializes the DocFileHandler object with the specified file path.
        :param file_path: str, the path to the Word document file.
        """
        self.file_path = file_path
        self.document = Document(self.file_path)

    def read_text(self):
        """
        Reads the content of a Word document and returns it as a string.
        :return: str, the content of the Word document.
        """
        return '\n'.join([paragraph.text for paragraph in self.document.paragraphs])

    def write_text(self, content, font_size=12, alignment='left'):
        """
        Writes the specified content to a Word document.
        :param content: str, the text content to write.
        :param font_size: int, optional, the font size of the text (default is 12).
        :param alignment: str, optional, the alignment of the text ('left', 'center', or 'right'; default is 'left').
        :return: bool, True if the write operation is successful, False otherwise.
        """
        paragraph = self.document.add_paragraph(content)
        if font_size > 0:
            paragraph.add_run().add_font_size(font_size)
        if alignment:
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT(alignment)
        return True

    def add_heading(self, heading, level=1):
        """
        Adds a heading to the Word document.
        :param heading: str, the text of the heading.
        :param level: int, optional, the level of the heading (1, 2, 3, etc.; default is 1).
        :return: bool, True if the heading is successfully added, False otherwise.
        """
        if level > 3:
            return False
        heading_style = self.document.styles['Normal']
        if not heading_style:
            heading_style = self.document.styles.add_style('Heading' + str(level), WD_PARAGRAPH_ALIGNMENT=WD_PARAGRAPH_ALIGNMENT['center'])
        self.document.add_heading(heading, style=heading_style)
        return True

    def add_table(self, data):
        """
        Adds a table to the Word document with the specified data.
        :param data: list of lists, the data to populate the table.
        :return: bool, True if the table is successfully added, False otherwise.
        """
        table = self.document.add_table(rows=1, cols=len(data[0]))
        for i, row in enumerate(data):
            for j, value in enumerate(row):
                table.cell(i).paragraphs[0].add_run(str(value))
        return True

    def _get_alignment_value(self, alignment):
        """
        Returns the alignment value corresponding to the given alignment string.
        :param alignment: str, the alignment string ('left', 'center', or 'right').
        :return: int, the alignment value.
        """
        return WD_PARAGRAPH_ALIGNMENT(alignment).value

if __name__ == "__main__":
    # Test case for read_text method
    handler = DocFileHandler('test.docx')
    output = handler.read_text()
    print(output)

    # Test case for write_text method
    handler = DocFileHandler('test.docx')
    handler.write_text("New Content!", font_size=14, alignment='center')
    output = handler.read_text()
    print(output)

    # Test case for add_heading method
    handler = DocFileHandler('test.docx')
    handler.add_heading("New Heading", level=2)
    output = handler.read_text()
    print(output)

    # Test case for add_table method
    handler = DocFileHandler('test.docx')
    data = [['Row1Cell1', 'Row1Cell2'], ['Row2Cell1', 'Row2Cell2']]
    handler.add_table(data)
    output = handler.read_text()
    print(output)