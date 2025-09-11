from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


class DocFileHandler:
    """
    This is a class that handles Word documents and provides functionalities for reading, writing, and modifying the content of Word documents.
    """

    def __init__(self, file_path):
        """
        Initializes the DocFileHandler object with the specified file path.
        :param file_path: str, the path to the Word document file.
        """
        self.file_path = file_path

    def read_text(self):
        """
        Reads the content of a Word document and returns it as a string.
        :return: str, the content of the Word document.
        """
        try:
            document = Document(self.file_path)
            text = ''
            for para in document.paragraphs:
                text += para.text + '\n'
            return text
        except Exception as e:
            print(f"Error reading file: {e}")
            return None

    def write_text(self, content, font_size=12, alignment='left'):
        """
        Writes the specified content to a Word document.
        :param content: str, the text content to write.
        :param font_size: int, optional, the font size of the text (default is 12).
        :param alignment: str, optional, the alignment of the text ('left', 'center', or 'right'; default is 'left').
        :return: bool, True if the write operation is successful, False otherwise.
        """
        try:
            document = Document()
            para = document.add_paragraph()
            run = para.add_run(content)
            font = run.font
            font.name = 'Calibri'
            font.size = Pt(font_size)
            alignment_value = self._get_alignment_value(alignment)
            para.paragraph_format.alignment = alignment_value
            document.save(self.file_path)
            return True
        except Exception as e:
            print(f"Error writing file: {e}")
            return False

    def add_heading(self, heading, level=1):
        """
        Adds a heading to the Word document.
        :param heading: str, the text of the heading.
        :param level: int, optional, the level of the heading (1, 2, 3, etc.; default is 1).
        :return: bool, True if the heading is successfully added, False otherwise.
        """
        try:
            document = Document(self.file_path)
            heading_para = document.add_heading(heading, level)
            document.save(self.file_path)
            return True
        except Exception as e:
            print(f"Error adding heading: {e}")
            return False

    def add_table(self, data):
        """
        Adds a table to the Word document with the specified data.
        :param data: list of lists, the data to populate the table.
        :return: bool, True if the table is successfully added, False otherwise.
        """
        try:
            document = Document()
            table = document.add_table(rows=len(data), cols=len(data[0]))
            for i in range(len(data)):
                for j in range(len(data[0])):
                    cell = table.cell(i, j)
                    cell.text = str(data[i][j])
            document.save(self.file_path)
            return True
        except Exception as e:
            print(f"Error adding table: {e}")
            return False

    def _get_alignment_value(self, alignment):
        """
        Returns the alignment value corresponding to the given alignment string.
        :param alignment: str, the alignment string ('left', 'center', or 'right').
        :return: int, the alignment value.
        """
        if alignment == 'left':
            return WD_PARAGRAPH_ALIGNMENT.LEFT
        elif alignment == 'center':
            return WD_PARAGRAPH_ALIGNMENT.CENTER
        elif alignment == 'right':
            return WD_PARAGRAPH_ALIGNMENT.RIGHT
        else:
            raise ValueError("Invalid alignment value")


if __name__ == "__main__":
    # Test case for read_text method
    handler = DocFileHandler('test.docx')
    output = handler.read_text()
    print(output)

    # Test case for write_text method
    handler = DocFileHandler('test.docx')
    content = "This is a sample text."
    output = handler.write_text(content)
    print(output)

    # Test case for add_heading method
    handler = DocFileHandler('test.docx')
    heading = "Sample Heading"
    output = handler.add_heading(heading)
    print(output)

    # Test case for add_table method
    handler = DocFileHandler('test.docx')
    data = [[1, 2, 3], [4, 5, 6], [7, 8, 9]]
    output = handler.add_table(data)
    print(output)